"""Word report generation.

Structure mirrors the Reddit pipeline's narrative-report style: prose first,
tables as supporting evidence, methodology and limitations stated explicitly at
the end rather than buried.

The methodology section is not boilerplate. It records sentiment coverage, the
provider used, and the scope limits of the data — because a report that says
"net sentiment -0.34" without saying "on 61% of rows, scored by a third-party
actor with known weakness on Hinglish" invites the reader to over-trust it.
"""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from igpulse.analyze.metrics import FigureMetrics, NarrativeMetrics
from igpulse.analyze.sentiment import SentimentSummary
from igpulse.analyze.themes import ThemeTerm
from igpulse.config import AppConfig

logger = logging.getLogger(__name__)

_ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def _fmt_pct(value: float | None, *, signed: bool = False) -> str:
    if value is None:
        return "n/a"
    return f"{value:+.1f}%" if signed else f"{value:.1f}%"


def _fmt_net(value: float | None) -> str:
    return "n/a" if value is None else f"{value:+.2f}"


# Usable width inside default 1" margins on A4: 6.27". Column widths must be
# set on EVERY cell, not just the column object — python-docx writes widths per
# cell, and Word ignores a column width that its cells contradict. Without this
# the table auto-fits and headers wrap mid-word ("Engageme / nt").
_TABLE_WIDTH_INCHES = 6.27


def _add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    weights: list[float] | None = None,
    font_size: int = 9,
):
    if weights is None:
        weights = [1.0] * len(headers)
    if len(weights) != len(headers):
        raise ValueError("weights must match header count")

    total = sum(weights)
    widths = [
        Inches(_TABLE_WIDTH_INCHES * w / total) for w in weights
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(font_size)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].width = widths[idx]
            cells[idx].text = value
            for para in cells[idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(font_size)

    doc.add_paragraph()
    return table


def build_report(
    cfg: AppConfig,
    *,
    narratives: list[NarrativeMetrics],
    figures: list[FigureMetrics],
    themes: list[ThemeTerm],
    sentiment: SentimentSummary,
    generated_at: datetime,
    output_path: Path | None = None,
) -> Path:
    out_dir = Path(cfg.settings.report.output_dir)
    out_dir.mkdir(parents=True, exist_ok=True)
    local_time = generated_at.astimezone(cfg.timezone)
    if output_path is None:
        output_path = out_dir / (
            f"{cfg.settings.project.client_label}_"
            f"{local_time:%Y%m%d_%H%M}_report.docx"
        )

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # -- title ------------------------------------------------------------ #
    title = doc.add_heading("Instagram Discourse Report", level=0)
    for run in title.runs:
        run.font.color.rgb = _ACCENT
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = subtitle.add_run(
        f"{cfg.settings.project.client_label} · "
        f"{local_time:%d %B %Y, %H:%M} "
        f"({cfg.settings.project.timezone})"
    )
    run.italic = True
    run.font.size = Pt(10)

    sections = cfg.settings.report.docx.include_sections

    # -- executive summary ------------------------------------------------ #
    doc.add_heading("Executive summary", level=1)
    if narratives:
        top = narratives[0]
        movers = [n for n in narratives if n.volume_delta_pct is not None]
        movers.sort(key=lambda n: abs(n.volume_delta_pct or 0), reverse=True)
        doc.add_paragraph(
            f"Across {len(narratives)} tracked narratives, "
            f"{sum(n.post_count for n in narratives):,} posts and "
            f"{sum(n.comment_count for n in narratives):,} comments were "
            f"collected. The highest-volume narrative was "
            f"“{top.label}” at {top.post_count:,} posts "
            f"({top.share_of_voice * 100:.1f}% share of voice)."
        )
        if movers:
            mover = movers[0]
            doc.add_paragraph(
                f"The sharpest movement was “{mover.label}”, "
                f"{_fmt_pct(mover.volume_delta_pct, signed=True)} in volume "
                f"against the preceding period."
            )
    else:
        doc.add_paragraph(
            "No narrative data was collected in this run. Check that "
            "config/narratives.yaml defines at least one narrative and that "
            "the lookback window covers recent activity."
        )

    # -- narrative section ------------------------------------------------ #
    if "narrative" in sections and narratives:
        doc.add_heading("Narrative landscape", level=1)
        doc.add_paragraph(
            "Volume, reach and audience reaction by issue. Engagement is "
            "reported in absolute terms because follower denominators are not "
            "collected for non-allowlisted accounts."
        )
        _add_table(
            doc,
            ["Narrative", "Posts", "Comments", "Engmt.",
             "Share", "Vol. Δ", "Net sent."],
            [
                [
                    n.label,
                    f"{n.post_count:,}",
                    f"{n.comment_count:,}",
                    f"{n.total_engagement:,}",
                    f"{n.share_of_voice * 100:.1f}%",
                    _fmt_pct(n.volume_delta_pct, signed=True),
                    _fmt_net(n.comment_sentiment.net_sentiment),
                ]
                for n in narratives
            ],
            weights=[4.4, 1.0, 1.2, 1.3, 1.0, 1.1, 1.2],
        )

        if themes:
            doc.add_heading("Language carrying each narrative", level=2)
            by_narrative: dict[str, list[ThemeTerm]] = {}
            for term in themes:
                by_narrative.setdefault(term.narrative_key, []).append(term)
            label_of = {n.narrative_key: n.label for n in narratives}
            for key, terms in by_narrative.items():
                para = doc.add_paragraph()
                para.add_run(f"{label_of.get(key, key)}: ").bold = True
                para.add_run(", ".join(t.term for t in terms[:12]))

    # -- public figure section -------------------------------------------- #
    if "public_figure" in sections:
        doc.add_heading("Public figures", level=1)
        tracked = [f for f in figures if f.category != "own_side"]
        if tracked:
            doc.add_paragraph(
                "Elected officials, party accounts and registered media "
                "organisations on the curated allowlist. Audience sentiment "
                "reflects comments on their posts, not their own posting."
            )
            _add_table(
                doc,
                ["Account", "Category", "Followers", "Posts",
                 "Avg engmt.", "Eng. rate", "Aud. net sent."],
                [
                    [
                        f.display_name,
                        f.category.replace("_", " "),
                        f"{f.follower_count:,}" if f.follower_count else "n/a",
                        f"{f.post_count:,}",
                        f"{f.avg_engagement_per_post:,.0f}",
                        _fmt_pct(f.engagement_rate * 100)
                        if f.engagement_rate is not None else "n/a",
                        _fmt_net(f.audience_sentiment.net_sentiment),
                    ]
                    for f in tracked
                ],
                weights=[2.6, 1.7, 1.4, 0.8, 1.3, 1.1, 1.4],
            )
        else:
            doc.add_paragraph(
                "No public figures configured. Add entries to "
                "config/public_figures.yaml."
            )

    # -- own side section -------------------------------------------------- #
    if "own_side" in sections:
        doc.add_heading("Own-side performance", level=1)
        own = [f for f in figures if f.category == "own_side"]
        if own:
            _add_table(
                doc,
                ["Account", "Followers", "Posts", "Total engmt.",
                 "Avg/post", "Eng. rate", "Aud. net sent."],
                [
                    [
                        f.display_name,
                        f"{f.follower_count:,}" if f.follower_count else "n/a",
                        f"{f.post_count:,}",
                        f"{f.total_engagement:,}",
                        f"{f.avg_engagement_per_post:,.0f}",
                        _fmt_pct(f.engagement_rate * 100)
                        if f.engagement_rate is not None else "n/a",
                        _fmt_net(f.audience_sentiment.net_sentiment),
                    ]
                    for f in own
                ],
                weights=[2.6, 1.4, 0.8, 1.4, 1.2, 1.1, 1.4],
            )
            if tracked_avg := _benchmark(figures):
                doc.add_paragraph(
                    f"Own-side average engagement per post is "
                    f"{tracked_avg[0]:,.0f} against {tracked_avg[1]:,.0f} for "
                    f"tracked public figures."
                )
        else:
            doc.add_paragraph(
                "No own-side accounts configured. Add entries with category "
                "'own_side' to config/public_figures.yaml."
            )

    # -- methodology ------------------------------------------------------- #
    if "methodology" in sections:
        doc.add_heading("Methodology and limitations", level=1)
        coverage = sentiment.coverage * 100
        for line in [
            f"Collection: Apify actors "
            f"{cfg.settings.apify.actors.instagram_scraper} and "
            f"{cfg.settings.apify.actors.instagram_comment_scraper}. "
            f"Lookback {cfg.settings.ingest.narrative.lookback} for the "
            f"narrative lens.",

            f"Sentiment: {cfg.settings.sentiment.actor_id} via Apify. "
            f"{sentiment.scored:,} rows scored with confidence at or above "
            f"{cfg.settings.sentiment.min_confidence:.2f}; "
            f"{sentiment.uncertain:,} fell below that threshold and are "
            f"recorded as uncertain rather than assigned a polarity. "
            f"Effective coverage {coverage:.1f}%.",

            "Known weakness: the sentiment actor is trained primarily on "
            "English. Code-mixed Hinglish and Devanagari text score less "
            "reliably, so sentiment figures should be read as directional "
            "rather than precise.",

            "Sampling: Instagram search returns a ranked subset, not a "
            "census. Volume figures are comparable between runs of this "
            "pipeline but are not estimates of total platform activity.",

            "Scope: accounts are named in this report only where they appear "
            "on the curated public-figure allowlist — elected officials, "
            "party accounts, registered media, and the client's own "
            "accounts. All other authors contribute to aggregate counts "
            "under per-run pseudonyms and are not tracked between runs.",
        ]:
            para = doc.add_paragraph(line, style="List Bullet")
            para.paragraph_format.space_after = Pt(6)

    doc.save(str(output_path))
    logger.info("wrote %s", output_path)
    return output_path


def _benchmark(figures: list[FigureMetrics]) -> tuple[float, float] | None:
    own = [f for f in figures if f.category == "own_side" and f.post_count]
    other = [f for f in figures if f.category != "own_side" and f.post_count]
    if not own or not other:
        return None
    return (
        sum(f.avg_engagement_per_post for f in own) / len(own),
        sum(f.avg_engagement_per_post for f in other) / len(other),
    )
